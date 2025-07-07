"""
Document processing module for extracting text and metadata from various file formats.
Supports PDF, DOCX, Markdown, and plain text files.
"""

import asyncio
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
import hashlib
import mimetypes
# Document processing libraries
import PyPDF2
import fitz  # PyMuPDF for advanced PDF processing
from docx import Document
import markdown
from bs4 import BeautifulSoup

# Text processing
import re
from langchain.text_splitter import RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)

class DocumentChunk:
    """Represents a chunk of processed document content."""
    
    def __init__(self, content: str, metadata: Dict[str, Any]):
        self.content = content
        self.metadata = metadata
        self.chunk_id = self._generate_id()
    
    def _generate_id(self) -> str:
        """Generate unique ID for this chunk."""
        content_hash = hashlib.md5(self.content.encode()).hexdigest()
        return f"chunk_{content_hash[:12]}"
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert chunk to dictionary for storage."""
        return {
            'chunk_id': self.chunk_id,
            'content': self.content,
            'metadata': self.metadata
        }

class DocumentProcessor:
    """Main document processing class."""
    
    def __init__(self, config):
        self.config = config
        self.text_splitter = None
        self.supported_processors = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.md': self._process_markdown,
            '.txt': self._process_text
        }
    
    async def initialize(self):
        """Initialize the document processor."""
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.config.chunk_size,
            chunk_overlap=self.config.chunk_overlap,
            separators=self.config.get('chunking.separators', ["\n\n", "\n", " ", ""])
        )
        logger.info("Document processor initialized")
    
    async def process_file(self, file_path: str) -> List[DocumentChunk]:
        """Process a single file and return chunks."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Check file size
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        max_size = self.config.get('processing.max_file_size_mb', 100)
        if file_size_mb > max_size:
            raise ValueError(f"File too large: {file_size_mb:.1f}MB > {max_size}MB")
        
        # Check file extension
        extension = file_path.suffix.lower()
        if extension not in self.config.supported_extensions:
            raise ValueError(f"Unsupported file type: {extension}")
        
        # Extract text and metadata
        text_content, metadata = await self._extract_content(file_path)
        
        # Create chunks
        chunks = await self._create_chunks(text_content, metadata)
        
        logger.info(f"Processed {file_path.name}: {len(chunks)} chunks created")
        return chunks
    
    async def _extract_content(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Extract text content and metadata from file."""
        extension = file_path.suffix.lower()
        processor = self.supported_processors.get(extension)
        
        if not processor:
            raise ValueError(f"No processor for file type: {extension}")
        
        # Base metadata
        metadata = {
            'source': str(file_path),
            'filename': file_path.name,
            'extension': extension,
            'file_size': file_path.stat().st_size,
            'modified_time': file_path.stat().st_mtime
        }
        
        # Extract content using appropriate processor
        content, extra_metadata = await processor(file_path)
        metadata.update(extra_metadata)
        
        return content, metadata
    
    async def _process_pdf(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Process PDF file and extract text with advanced highlight detection."""
        content = ""
        metadata = {
            'pages': 0,
            'highlights': [],
            'annotations': [],
            'highlight_count': 0,
            'annotation_count': 0
        }

        try:
            # Use PyMuPDF for advanced PDF processing
            pdf_document = fitz.open(str(file_path))
            metadata['pages'] = len(pdf_document)

            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]

                # Extract text
                page_text = page.get_text()
                content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"

                # Extract highlights and annotations
                if self.config.get('processing.extract_highlights', True):
                    page_highlights, page_annotations = self._extract_pdf_annotations(page, page_num + 1)
                    metadata['highlights'].extend(page_highlights)
                    metadata['annotations'].extend(page_annotations)

            # Update counts
            metadata['highlight_count'] = len(metadata['highlights'])
            metadata['annotation_count'] = len(metadata['annotations'])

            # Extract PDF metadata
            pdf_metadata = pdf_document.metadata
            if pdf_metadata:
                metadata.update({
                    'title': pdf_metadata.get('title', ''),
                    'author': pdf_metadata.get('author', ''),
                    'subject': pdf_metadata.get('subject', ''),
                    'creator': pdf_metadata.get('creator', ''),
                    'producer': pdf_metadata.get('producer', ''),
                    'creation_date': pdf_metadata.get('creationDate', ''),
                    'modification_date': pdf_metadata.get('modDate', '')
                })

            pdf_document.close()

        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}")
            # Fallback to PyPDF2 if PyMuPDF fails
            try:
                content, fallback_metadata = await self._process_pdf_fallback(file_path)
                metadata.update(fallback_metadata)
            except Exception as fallback_error:
                logger.error(f"Fallback PDF processing also failed: {fallback_error}")
                raise e

        return content.strip(), metadata
    
    def _extract_pdf_annotations(self, page, page_num: int) -> tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        """Extract highlights and annotations from PDF page using PyMuPDF."""
        highlights = []
        annotations = []

        try:
            # Get all annotations on the page
            annot_list = page.annots()

            for annot in annot_list:
                annot_dict = annot.info
                annot_type = annot_dict.get('type', '')

                # Extract annotation content
                content = annot_dict.get('content', '').strip()

                # Get annotation rectangle (position)
                rect = annot.rect

                # Get annotation color if available
                color = None
                if hasattr(annot, 'colors') and annot.colors:
                    color = annot.colors.get('stroke', annot.colors.get('fill'))

                # Base annotation data
                base_data = {
                    'page': page_num,
                    'type': annot_type,
                    'content': content,
                    'position': {
                        'x0': rect.x0,
                        'y0': rect.y0,
                        'x1': rect.x1,
                        'y1': rect.y1
                    },
                    'color': self._format_color(color) if color else None,
                    'author': annot_dict.get('title', ''),
                    'creation_date': annot_dict.get('creationDate', ''),
                    'modification_date': annot_dict.get('modDate', '')
                }

                # Handle different annotation types
                if annot_type in ['Highlight', 'Squiggly', 'StrikeOut', 'Underline']:
                    # These are highlight-type annotations
                    highlight_data = base_data.copy()

                    # Try to extract the highlighted text
                    highlighted_text = self._extract_highlighted_text(page, rect)
                    if highlighted_text:
                        highlight_data['highlighted_text'] = highlighted_text
                        highlight_data['text_length'] = len(highlighted_text)

                    # Determine highlight color category
                    highlight_data['color_category'] = self._categorize_highlight_color(color)

                    highlights.append(highlight_data)

                elif annot_type in ['Text', 'FreeText', 'Note']:
                    # These are note/comment annotations
                    note_data = base_data.copy()

                    # Try to get associated text
                    if not content:
                        # Try to extract text near the annotation
                        nearby_text = self._extract_nearby_text(page, rect)
                        if nearby_text:
                            note_data['nearby_text'] = nearby_text

                    annotations.append(note_data)

                else:
                    # Other annotation types (stamps, drawings, etc.)
                    annotations.append(base_data)

        except Exception as e:
            logger.warning(f"Error extracting annotations from page {page_num}: {e}")

        return highlights, annotations

    def _extract_highlighted_text(self, page, rect) -> str:
        """Extract text that is highlighted within the given rectangle."""
        try:
            # Get text blocks that intersect with the highlight rectangle
            text_blocks = page.get_text("dict")
            highlighted_text = ""

            for block in text_blocks.get("blocks", []):
                if "lines" in block:
                    for line in block["lines"]:
                        line_rect = fitz.Rect(line["bbox"])
                        # Check if line intersects with highlight rectangle
                        if line_rect.intersects(rect):
                            for span in line.get("spans", []):
                                span_rect = fitz.Rect(span["bbox"])
                                if span_rect.intersects(rect):
                                    highlighted_text += span.get("text", "")

            return highlighted_text.strip()
        except Exception as e:
            logger.warning(f"Error extracting highlighted text: {e}")
            return ""

    def _extract_nearby_text(self, page, rect, margin=50) -> str:
        """Extract text near an annotation for context."""
        try:
            # Expand rectangle to get nearby text
            expanded_rect = fitz.Rect(
                rect.x0 - margin,
                rect.y0 - margin,
                rect.x1 + margin,
                rect.y1 + margin
            )

            # Get text in the expanded area
            nearby_text = page.get_textbox(expanded_rect)
            return nearby_text.strip()
        except Exception as e:
            logger.warning(f"Error extracting nearby text: {e}")
            return ""

    def _format_color(self, color) -> str:
        """Format color information to a readable string."""
        try:
            if isinstance(color, (list, tuple)) and len(color) >= 3:
                # RGB color
                r, g, b = color[:3]
                return f"rgb({int(r*255)}, {int(g*255)}, {int(b*255)})"
            elif isinstance(color, dict):
                # Color dictionary
                if 'rgb' in color:
                    rgb = color['rgb']
                    return f"rgb({int(rgb[0]*255)}, {int(rgb[1]*255)}, {int(rgb[2]*255)})"
            return str(color)
        except:
            return "unknown"

    def _categorize_highlight_color(self, color) -> str:
        """Categorize highlight color into common categories."""
        if not color:
            return "default"

        try:
            color_str = str(color).lower()
            if "yellow" in color_str or "255, 255, 0" in color_str:
                return "yellow"
            elif "red" in color_str or "255, 0, 0" in color_str:
                return "red"
            elif "green" in color_str or "0, 255, 0" in color_str:
                return "green"
            elif "blue" in color_str or "0, 0, 255" in color_str:
                return "blue"
            elif "orange" in color_str or "255, 165, 0" in color_str:
                return "orange"
            elif "pink" in color_str or "255, 192, 203" in color_str:
                return "pink"
            else:
                return "other"
        except:
            return "unknown"

    async def _process_pdf_fallback(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Fallback PDF processing using PyPDF2 when PyMuPDF fails."""
        content = ""
        metadata = {'pages': 0, 'highlights': [], 'annotations': []}

        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata['pages'] = len(pdf_reader.pages)

                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"

                # Extract basic PDF metadata
                if pdf_reader.metadata:
                    metadata.update({
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', '')
                    })

        except Exception as e:
            logger.error(f"Fallback PDF processing failed for {file_path}: {str(e)}")
            raise

        return content.strip(), metadata
    
    async def _process_docx(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Process DOCX file and extract text."""
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            content = ""
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # Extract metadata
            metadata = {
                'paragraphs': len(doc.paragraphs),
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'subject': doc.core_properties.subject or '',
                'created': str(doc.core_properties.created) if doc.core_properties.created else ''
            }
            
            return content.strip(), metadata
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            raise
    
    async def _process_markdown(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Process Markdown file and extract text."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            # Convert markdown to HTML then extract text
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, 'html.parser')
            text_content = soup.get_text()
            
            # Extract markdown-specific metadata
            metadata = {
                'format': 'markdown',
                'headers': self._extract_markdown_headers(md_content),
                'links': self._extract_markdown_links(md_content)
            }
            
            return text_content.strip(), metadata
            
        except Exception as e:
            logger.error(f"Error processing Markdown {file_path}: {str(e)}")
            raise
    
    async def _process_text(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Process plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            metadata = {
                'format': 'text',
                'lines': len(content.splitlines()),
                'words': len(content.split())
            }
            
            return content.strip(), metadata
            
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {str(e)}")
            raise
    
    def _extract_markdown_headers(self, content: str) -> List[Dict[str, Any]]:
        """Extract headers from markdown content."""
        headers = []
        for match in re.finditer(r'^(#{1,6})\s+(.+)$', content, re.MULTILINE):
            headers.append({
                'level': len(match.group(1)),
                'text': match.group(2).strip()
            })
        return headers
    
    def _extract_markdown_links(self, content: str) -> List[Dict[str, Any]]:
        """Extract links from markdown content."""
        links = []
        for match in re.finditer(r'\[([^\]]+)\]$$([^)]+)$$', content):
            links.append({
                'text': match.group(1),
                'url': match.group(2)
            })
        return links
    
    async def _create_chunks(self, content: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Split content into chunks."""
        if not content.strip():
            return []
        # Split text into chunks
        text_chunks = self.text_splitter.split_text(content)
        # Create DocumentChunk objects
        chunks = []
        for i, chunk_text in enumerate(text_chunks):
            chunk_metadata = metadata.copy()
            chunk_metadata.update({
                'chunk_index': i,
                'total_chunks': len(text_chunks),
                'chunk_size': len(chunk_text)
            })
            chunk = DocumentChunk(chunk_text, chunk_metadata)
            chunks.append(chunk)
        return chunks
